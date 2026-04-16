import json
import os
from datetime import datetime

from flask import Blueprint, render_template, redirect, url_for, request, flash, Response, stream_with_context, current_app
from flask_login import login_required, current_user

from app.models import FigurinoSheet, EventRole
from app.constants import RoleName
from app.storage import save_file, delete_file
from .. import db

figurino_bp = Blueprint("figurino", __name__)

_ALLOWED_PHOTO_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}


def _parse_pieces(form) -> list:
    """Parse piece_names[] + piece_qtys[] from form into [{"name", "qty"}] list."""
    names = form.getlist("piece_names[]")
    qtys  = form.getlist("piece_qtys[]")
    result = []
    for name, qty_str in zip(names, qtys):
        name = name.strip()
        if name:
            try:
                qty = max(1, int(qty_str or "1"))
            except ValueError:
                qty = 1
            result.append({"name": name, "qty": qty})
    return result


# ── Listing ────────────────────────────────────────────────────

@figurino_bp.route("/figurinos")
@login_required
def figurinos():
    from .drive_service import normalize_name

    sheets = FigurinoSheet.query.order_by(FigurinoSheet.character_name.asc()).all()
    sheet_norms = {s.character_name_norm for s in sheets if s.character_name_norm}

    all_chars = db.session.query(EventRole.character_name).distinct().all()
    chars_without_sheet = sorted(
        {c[0] for c in all_chars if c[0] and normalize_name(c[0]) not in sheet_norms}
    )

    return render_template(
        "figurinos.html",
        sheets=sheets,
        chars_without_sheet=chars_without_sheet,
    )


# ── Create ─────────────────────────────────────────────────────

@figurino_bp.route("/figurinos/new", methods=["GET", "POST"])
@login_required
def new_sheet():
    from .drive_service import normalize_name

    if request.method == "POST":
        character_name = request.form.get("character_name", "").strip()
        if not character_name:
            flash("Nome do personagem é obrigatório.")
            return render_template("figurino_form.html", sheet=None,
                                   pieces=_parse_pieces(request.form) or [{"name": "", "qty": 1}],
                                   title="Nova Ficha de Figurino")

        pieces = _parse_pieces(request.form)
        notes = request.form.get("notes", "").strip()

        photo_url = None
        photo_file = request.files.get("photo")
        if photo_file and photo_file.filename:
            ext = os.path.splitext(photo_file.filename)[1].lower()
            if ext in _ALLOWED_PHOTO_EXTENSIONS:
                photo_url = save_file(photo_file, "figurino_photos")

        sheet = FigurinoSheet(
            character_name=character_name,
            character_name_norm=normalize_name(character_name),
            photo_filename=photo_url,
            pieces=json.dumps(pieces, ensure_ascii=False) if pieces else None,
            notes=notes or None,
        )
        db.session.add(sheet)
        from app.utils import audit
        audit("create", "figurino", None, character_name, "Ficha de figurino criada")
        db.session.commit()
        flash(f'Ficha de "{character_name}" criada com sucesso!')
        return redirect(url_for("figurino.figurinos"))

    return render_template("figurino_form.html", sheet=None,
                           pieces=[{"name": "", "qty": 1}],
                           title="Nova Ficha de Figurino")


# ── Edit ───────────────────────────────────────────────────────

@figurino_bp.route("/figurinos/<int:sheet_id>/edit", methods=["GET", "POST"])
@login_required
def edit_sheet(sheet_id: int):
    from .drive_service import normalize_name

    sheet = FigurinoSheet.query.get_or_404(sheet_id)

    if request.method == "POST":
        character_name = request.form.get("character_name", "").strip()
        if not character_name:
            flash("Nome do personagem é obrigatório.")
            return render_template("figurino_form.html", sheet=sheet,
                                   pieces=sheet.pieces_list or [{"name": "", "qty": 1}],
                                   title=f"Editar — {sheet.character_name}")

        pieces = _parse_pieces(request.form)
        notes = request.form.get("notes", "").strip()

        photo_file = request.files.get("photo")
        if photo_file and photo_file.filename:
            ext = os.path.splitext(photo_file.filename)[1].lower()
            if ext in _ALLOWED_PHOTO_EXTENSIONS:
                delete_file(sheet.photo_filename)
                sheet.photo_filename = save_file(photo_file, "figurino_photos")

        sheet.character_name = character_name
        sheet.character_name_norm = normalize_name(character_name)
        sheet.pieces = json.dumps(pieces, ensure_ascii=False) if pieces else None
        sheet.notes = notes or None
        sheet.updated_at = datetime.utcnow()
        from app.utils import audit
        audit("edit", "figurino", sheet.id, character_name, "Ficha de figurino editada")
        db.session.commit()
        flash(f'Ficha de "{character_name}" atualizada!')
        return redirect(url_for("figurino.figurinos"))

    return render_template("figurino_form.html", sheet=sheet,
                           pieces=sheet.pieces_list or [{"name": "", "qty": 1}],
                           title=f"Editar — {sheet.character_name}")


# ── Print: single sheet ────────────────────────────────────────

@figurino_bp.route("/figurinos/<int:sheet_id>/print")
@login_required
def print_sheet(sheet_id: int):
    sheet = FigurinoSheet.query.get_or_404(sheet_id)
    items = [{"sheet": sheet, "pieces": sheet.pieces_list, "role": None, "talent": None}]
    return render_template("figurino_print.html", items=items, event=None,
                           title=f"Ficha: {sheet.character_name}")


# ── Print: all sheets for an event ────────────────────────────

@figurino_bp.route("/figurinos/print-event/<int:event_id>")
@login_required
def print_event_figurinos(event_id: int):
    from app.models import CalendarEvent
    from .drive_service import normalize_name

    event = CalendarEvent.query.get_or_404(event_id)

    items = []
    for role in event.roles:
        sheet = role.figurino_sheet
        if not sheet:
            norm = normalize_name(role.character_name)
            sheet = FigurinoSheet.query.filter_by(character_name_norm=norm).first()

        items.append({
            "role": role,
            "sheet": sheet,
            "pieces": sheet.pieces_list if sheet else [],
            "talent": role.talent,
        })

    return render_template("figurino_print.html", items=items, event=event,
                           title=f"Figurinos — {event.title}")


# ── Delete ─────────────────────────────────────────────────────

@figurino_bp.route("/figurinos/<int:sheet_id>/rotate-photo", methods=["POST"])
@login_required
def rotate_photo(sheet_id: int):
    sheet = FigurinoSheet.query.get_or_404(sheet_id)

    if not sheet.photo_filename:
        flash("Sem foto para girar.")
        return redirect(url_for("figurino.edit_sheet", sheet_id=sheet_id))

    photo_url = sheet.photo_filename
    if not photo_url.startswith("/uploads/"):
        flash("Formato de foto não suportado para rotação.")
        return redirect(url_for("figurino.edit_sheet", sheet_id=sheet_id))

    rel_path = photo_url[len("/uploads/"):]
    abs_path = os.path.join(current_app.config["UPLOAD_FOLDER"], rel_path)

    direction = request.form.get("direction", "cw")  # cw | ccw

    try:
        from PIL import Image
        img = Image.open(abs_path)
        degrees = -90 if direction == "cw" else 90
        img = img.rotate(degrees, expand=True)
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        img.save(abs_path, format="JPEG", quality=92, subsampling=0)
        sheet.updated_at = datetime.utcnow()
        db.session.commit()
    except Exception as e:
        flash(f"Erro ao girar foto: {e}")

    return redirect(url_for("figurino.edit_sheet", sheet_id=sheet_id))


@figurino_bp.route("/figurinos/<int:sheet_id>/delete", methods=["POST"])
@login_required
def delete_sheet(sheet_id: int):
    sheet = FigurinoSheet.query.get_or_404(sheet_id)

    delete_file(sheet.photo_filename)

    from app.utils import audit
    audit("delete", "figurino", sheet.id, sheet.character_name, "Ficha de figurino removida")
    EventRole.query.filter_by(figurino_sheet_id=sheet_id).update({"figurino_sheet_id": None})
    db.session.delete(sheet)
    db.session.commit()
    flash("Ficha removida do catálogo.")
    return redirect(url_for("figurino.figurinos"))


# ── Sync Drive (SUPERADMIN only) ───────────────────────────────

def _is_superadmin() -> bool:
    return current_user.is_authenticated and any(
        r.name == RoleName.SUPERADMIN for r in current_user.roles
    )


def _sync_normalize(text: str) -> str:
    import re as _re, unicodedata as _ud
    t = (text or "").strip().lower()
    t = _ud.normalize("NFKD", t).encode("ascii", "ignore").decode("ascii")
    t = _re.sub(r"[^a-z0-9]+", " ", t)
    return _re.sub(r"\s+", " ", t).strip()


def _sync_extract_name(doc) -> str:
    skip = ["coordenador", "data de", "horário", "durante",
            "item presente", "item ausente", "montagem", "retirada",
            "devolução", "manto produções"]

    def _is_name_para(para) -> str | None:
        text = para.text.strip()
        if not text or any(p in text.lower() for p in skip):
            return None
        if (text == text.upper() and len(text) > 3) or any(
            run.bold for run in para.runs if run.text.strip()
        ):
            return text.title()
        return None

    # 1. Parágrafos de nível raiz
    for para in doc.paragraphs:
        name = _is_name_para(para)
        if name:
            return name

    # 2. Células das primeiras 3 linhas de cada tabela (nome pode estar em header table)
    for table in doc.tables:
        for row in table.rows[:3]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    name = _is_name_para(para)
                    if name:
                        return name
    return ""


def _sync_extract_pieces(doc) -> list:
    """Encontra a tabela de peças (M/R/D/Qtd/Item = 5 colunas) e extrai os itens."""
    import re as _re

    def _try_table(table) -> list:
        pieces = []
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue  # pula cabeçalho
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 5 or not cells[4]:
                continue
            try:
                qty = int(_re.sub(r"\D", "", cells[3])) if cells[3] else 1
            except ValueError:
                qty = 1
            pieces.append({"name": cells[4], "qty": qty})
        return pieces

    for table in doc.tables:
        # Tabela de peças tem pelo menos 5 colunas e cabeçalho com "Item"
        if not table.rows:
            continue
        header_cells = [c.text.strip().upper() for c in table.rows[0].cells]
        if len(header_cells) >= 5 and any("ITEM" in h for h in header_cells):
            pieces = _try_table(table)
            if pieces:
                return pieces

    # Fallback: tenta todas as tabelas com ≥5 colunas
    for table in doc.tables:
        if table.rows and len(table.rows[0].cells) >= 5:
            pieces = _try_table(table)
            if pieces:
                return pieces

    return []


def _sync_save_photo(doc, file_id: str) -> tuple[str | None, str | None]:
    """Extrai a primeira imagem do doc e salva localmente.

    Returns:
        (url, error_msg) — url é None se não encontrou foto ou falhou;
        error_msg é None em caso de sucesso.
    """
    import io as _io
    if current_app.config.get("USE_S3"):
        return None, None
    upload_folder = current_app.config["UPLOAD_FOLDER"]
    fname = f"drive_{file_id}.jpg"
    save_dir = os.path.join(upload_folder, "figurino_photos")
    save_path = os.path.join(save_dir, fname)
    try:
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            if getattr(rel, "is_external", False):
                continue  # imagem linkada externamente, sem blob
            img_data = rel.target_part.blob
            os.makedirs(save_dir, exist_ok=True)
            try:
                from PIL import Image, ImageOps
                img = Image.open(_io.BytesIO(img_data))
                img = ImageOps.exif_transpose(img)
                # JPEG não suporta alpha — converte RGBA/LA/P para RGB
                if img.mode in ("RGBA", "LA", "P"):
                    img = img.convert("RGB")
                img.save(save_path, format="JPEG", quality=92, subsampling=0)
            except ImportError:
                with open(save_path, "wb") as f:
                    f.write(img_data)
            return f"/uploads/figurino_photos/{fname}", None
    except Exception as e:
        return None, str(e)
    return None, None  # nenhuma imagem no doc


@figurino_bp.route("/figurinos/sync-drive")
@login_required
def sync_drive_page():
    if not _is_superadmin():
        flash("Acesso restrito.", "error")
        return redirect(url_for("figurino.figurinos"))
    return render_template("figurino_sync.html")


@figurino_bp.route("/figurinos/sync-drive/stream")
@login_required
def sync_drive_stream():
    if not _is_superadmin():
        return Response("data: {}\n\n", content_type="text/event-stream", status=403)

    def generate():
        import json as _json
        import io as _io
        import re as _re

        GDOC_MIME = "application/vnd.google-apps.document"
        DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        FOLDER_ID = current_app.config.get(
            "FIGURINO_DRIVE_FOLDER_ID", "1XV_e6SprlXujo5Lav4lxAjvsVHhPSR-8"
        )

        def sse(data: dict) -> str:
            return f"data: {_json.dumps(data, ensure_ascii=False)}\n\n"

        env_json = os.environ.get("GOOGLE_SHEETS_CREDENTIALS")
        creds_path = os.path.join("instance", "credentials", "sheets_service_account.json")

        if not env_json and not os.path.exists(creds_path):
            yield sse({"type": "error", "msg": "Credenciais não encontradas. Configure GOOGLE_SHEETS_CREDENTIALS."})
            return

        try:
            yield sse({"type": "info", "msg": "Conectando ao Google Drive..."})
            from google.oauth2 import service_account
            from googleapiclient.discovery import build
            from googleapiclient.http import MediaIoBaseDownload
            from docx import Document

            scopes = ["https://www.googleapis.com/auth/drive.readonly"]
            if env_json:
                creds = service_account.Credentials.from_service_account_info(
                    _json.loads(env_json), scopes=scopes
                )
            else:
                creds = service_account.Credentials.from_service_account_file(
                    creds_path, scopes=scopes
                )
            service = build("drive", "v3", credentials=creds, cache_discovery=False)

            yield sse({"type": "info", "msg": "Listando arquivos..."})
            files, page_token = [], None
            while True:
                resp = service.files().list(
                    q=f"'{FOLDER_ID}' in parents and trashed=false",
                    fields="nextPageToken, files(id, name, webViewLink, mimeType)",
                    pageSize=200,
                    pageToken=page_token,
                ).execute()
                files.extend(resp.get("files", []))
                page_token = resp.get("nextPageToken")
                if not page_token:
                    break

            supported = [f for f in files if f.get("mimeType") in (GDOC_MIME, DOCX_MIME)]
            skipped_fmt = len(files) - len(supported)
            total = len(supported)
            if skipped_fmt:
                skipped_names = [f["name"] for f in files if f.get("mimeType") not in (GDOC_MIME, DOCX_MIME)]
                yield sse({"type": "info", "msg": f"{len(files)} itens na pasta — {total} Google Docs/DOCX para importar, {skipped_fmt} ignorado(s) por formato: {', '.join(skipped_names)}"})
            else:
                yield sse({"type": "info", "msg": f"{total} Google Docs encontrado(s). Importando..."})

            counts = {"created": 0, "updated": 0, "error": 0}

            for idx, file_info in enumerate(supported, 1):
                file_id   = file_info["id"]
                file_name = file_info["name"]
                mime_type = file_info.get("mimeType", GDOC_MIME)

                yield sse({"type": "progress", "idx": idx, "total": total, "name": file_name})

                try:
                    buf = _io.BytesIO()
                    req = (
                        service.files().export_media(fileId=file_id, mimeType=DOCX_MIME)
                        if mime_type == GDOC_MIME
                        else service.files().get_media(fileId=file_id)
                    )
                    downloader = MediaIoBaseDownload(buf, req)
                    done = False
                    while not done:
                        _, done = downloader.next_chunk()

                    doc = Document(_io.BytesIO(buf.getvalue()))
                    char_name = _sync_extract_name(doc) or _re.sub(
                        r"\.(docx?|gdoc)$", "", file_name, flags=_re.IGNORECASE
                    ).strip()
                    pieces    = _sync_extract_pieces(doc)
                    char_norm = _sync_normalize(char_name)
                    photo_url, photo_err = _sync_save_photo(doc, file_id)

                    sheet = FigurinoSheet.query.filter_by(drive_file_id=file_id).first()
                    if not sheet:
                        sheet = FigurinoSheet.query.filter_by(character_name_norm=char_norm).first()

                    action = "updated" if sheet else "created"
                    if not sheet:
                        sheet = FigurinoSheet()
                        db.session.add(sheet)

                    sheet.character_name      = char_name
                    sheet.character_name_norm = char_norm
                    sheet.pieces              = _json.dumps(pieces, ensure_ascii=False) if pieces else None
                    sheet.drive_file_id       = file_id
                    sheet.drive_url           = file_info.get("webViewLink", "")
                    sheet.last_synced_at      = datetime.utcnow()
                    sheet.updated_at          = datetime.utcnow()
                    if photo_url:
                        sheet.photo_filename = photo_url
                    db.session.commit()

                    counts[action] += 1
                    yield sse({"type": "result", "status": action,
                               "name": char_name, "pieces": len(pieces),
                               "photo": bool(photo_url),
                               "photo_err": photo_err})

                except Exception as e:
                    counts["error"] += 1
                    db.session.rollback()
                    yield sse({"type": "result", "status": "error",
                               "name": file_name, "msg": str(e)})

            # Diagnóstico: quais arquivos do Drive ainda não têm ficha?
            drive_ids = {f["id"] for f in supported}
            imported_ids = {
                s.drive_file_id
                for s in FigurinoSheet.query.filter(
                    FigurinoSheet.drive_file_id.in_(list(drive_ids))
                ).all()
            }
            missing_ids = drive_ids - imported_ids
            if missing_ids:
                missing_names = [
                    f["name"] for f in supported if f["id"] in missing_ids
                ]
                yield sse({"type": "missing", "files": missing_names})

            yield sse({"type": "done", **counts})

        except Exception as e:
            yield sse({"type": "error", "msg": f"Erro ao conectar: {e}"})

    return Response(
        stream_with_context(generate()),
        content_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
